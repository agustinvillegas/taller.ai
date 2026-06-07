from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import os


def agregar_indice(doc):
    p = doc.add_paragraph()
    run = p.add_run()

    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h'

    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')

    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def agregar_secciones(doc, secciones):
    for sec in secciones:
        titulo = sec.get("titulo", "")
        nivel = sec.get("nivel", 1)
        doc.add_heading(titulo, level=nivel)

        contenido = sec.get("contenido", "")
        if contenido:
            doc.add_paragraph(contenido)

        for sub in sec.get("subsecciones", []):
            agregar_secciones(doc, [sub])


def _agregar_pie_pagina(doc, texto: str):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(texto)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def generar_word(data: dict, output_path: str, empresa: dict = None):
    titulo = data.get("titulo", "Documento")
    terminos = data.get("terminos", [])
    secciones = data.get("secciones", [])
    imagenes_spec = data.get("imagenes", [])

    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(3)

    if empresa and empresa.get("nombre"):
        p_empresa = doc.add_paragraph()
        p_empresa.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_emp = p_empresa.add_run(empresa["nombre"].upper())
        run_emp.bold = True
        run_emp.font.size = Pt(20)
        run_emp.font.name = "Arial"
        run_emp.font.color.rgb = RGBColor(0x2F, 0x4F, 0x7F)

        if empresa.get("rubro"):
            p_rubro = doc.add_paragraph()
            p_rubro.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_rubro = p_rubro.add_run(empresa["rubro"].capitalize())
            run_rubro.font.size = Pt(13)
            run_rubro.font.name = "Arial"
            run_rubro.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        p_linea = doc.add_paragraph()
        p_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_linea = p_linea.add_run("─" * 40)
        run_linea.font.color.rgb = RGBColor(0x2F, 0x4F, 0x7F)

        datos_extra = []
        if empresa.get("cuit"):
            datos_extra.append(f"CUIT: {empresa['cuit']}")
        if empresa.get("direccion"):
            datos_extra.append(empresa["direccion"])
        if empresa.get("telefono"):
            datos_extra.append(f"Tel: {empresa['telefono']}")
        if datos_extra:
            p_datos = doc.add_paragraph()
            p_datos.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_datos = p_datos.add_run("  |  ".join(datos_extra))
            run_datos.font.size = Pt(10)
            run_datos.font.name = "Arial"
            run_datos.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        p_fecha = doc.add_paragraph()
        p_fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fecha_str = datetime.now().strftime("%d de %B de %Y")
        run_fecha = p_fecha.add_run(fecha_str)
        run_fecha.font.size = Pt(10)
        run_fecha.font.name = "Arial"
        run_fecha.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

        doc.add_paragraph()

        _agregar_pie_pagina(doc, f"{empresa['nombre']} — {datetime.now().strftime('%Y')}")

    titulo_par = doc.add_paragraph()
    titulo_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_par.paragraph_format.space_after = Pt(24)
    run = titulo_par.add_run(titulo)
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = "Arial"

    entries = data.get("indice", [])
    if entries:
        doc.add_heading("Indice", level=1)
        for e in entries:
            nivel = e.get("nivel", 1)
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Cm(0.5 * (nivel - 1))
            par.paragraph_format.space_after = Pt(2)
            par.paragraph_format.space_before = Pt(2)
            run = par.add_run(e.get("titulo", ""))
            run.font.size = Pt(11)
            if nivel == 1:
                run.bold = True
        doc.add_paragraph()
        agregar_indice(doc)
        doc.add_page_break()

    if secciones:
        agregar_secciones(doc, secciones)

    if terminos:
        for t in terminos:
            nombre = t.get("nombre", "")
            definicion = t.get("definicion", "")
            palabras_clave = [x.lower() for x in t.get("palabras_clave", [])]

            par = doc.add_paragraph()
            par.paragraph_format.space_before = Pt(10)
            par.paragraph_format.space_after = Pt(10)

            titulo_run = par.add_run(nombre + ": ")
            titulo_run.bold = True
            titulo_run.font.size = Pt(12)
            titulo_run.font.name = "Arial"

            restante = definicion
            while restante:
                posicion = len(restante)
                palabra = None

                for pk in palabras_clave:
                    idx = restante.lower().find(pk)
                    if idx != -1 and idx < posicion:
                        posicion = idx
                        palabra = pk

                if palabra is None:
                    run = par.add_run(restante)
                    run.font.size = Pt(12)
                    break

                if posicion > 0:
                    run = par.add_run(restante[:posicion])
                    run.font.size = Pt(12)

                sub = par.add_run(restante[posicion:posicion + len(palabra)])
                sub.underline = True
                sub.font.size = Pt(12)

                restante = restante[posicion + len(palabra):]

    if imagenes_spec:
        doc.add_page_break()

        for img_spec in imagenes_spec:
            if isinstance(img_spec, dict):
                img_path = img_spec.get("path", "")
                escala = img_spec.get("escala", 12)
            else:
                img_path = img_spec
                escala = 12

            if os.path.exists(img_path):
                try:
                    escala = max(5, min(18, float(escala)))
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_picture(img_path, width=Cm(escala))
                except Exception as e:
                    print("Error imagen:", e)
            else:
                print("Imagen no encontrada:", img_path)

    doc.save(output_path)
    print("Word creado:", output_path)
